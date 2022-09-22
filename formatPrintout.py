from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import subprocess
import json
import datetime
import random


def testPrint():
    print("Connected to formatter")

correct_answers = {
    "a1" : "2",
    "a2" : "37",
    "a3" : "2"
}


count =0


answer_lookup = {
    "en": {
        "a1": {
            "1": "do",
            "2": "do not",
        },
        "a3": {
            "1": "how you see and identify yourself",
            "2": "how others see and identify you",
            "3": "how you see and identify yourself and how others see and identify you", 
            "4": "neither how you see and identify yourself nor how others see and identify you"
        },
        "a2": {
            "0": "are",
            "1": "are",
            "2": "are",
            "3": "are",
            "4": "are",
            "5": "are",
            "6": "are",
            "7": "are",
            "8": "are",
            "9": "are", 
            "10": "are",
            "11": "are",
            "12": "are",
            "13": "are",
            "14": "are",
            "15": "are", 
            "16": "are",
            "17": "are",
            "18": "are",
            "19": "are",
            "20": "are", 
            "21": "are", 
            "22": "are",
            "23": "are", 
            "24": "are", 
            "25": "are", 
            "26": "are", 
            "27": "are", 
            "28": "are", 
            "29": "are",
            "30": "are", 
            "31": "are", 
            "32": "are", 
            "33": "are",
            "34": "are", 
            "35": "are", 
            "36": "are", 
            "37": "are not" 
        },
    },
    "es": {
        "a1": {
            "1": "Si",
            "2": "No"
        },
        "a3": {
            "1": "Cómo te ves e identificas",
            "2": "Como otres te ven e identifican",
            "3": "Todo lo anterior",
            "4": "Ninguna de las anteriores"
        },
               
        "a2": {
            "0": "Anderson",
            "1": "Barret",
            "2": "Buck",
            "3": "Bell",
            "4": "Coats",
            "5": "Cary",
            "6": "Kerry",
            "7": "Duperton",
            "8": "Dipitón", 
            "9": "Dishmey", 
            "10": "Dishmer",
            "11": "Buck",
            "12": "Ejice",
            "13": "Fuchue",
            "14": "Forche",
            "15": "Green", 
            "16": "Hilton", 
            "17": "Henderson",
            "18": "Hopkins",
            "19": "James",
            "20": "Jhonson", 
            "21": "Jones", 
            "22": "King", 
            "23": "Kinxon", 
            "24": "Miller", 
            "25": "Nwes", 
            "26": "Punez", 
            "27": "Paul", 
            "28": "Robinson", 
            "29": "Relmon", 
            "30": "Schod", 
            "31": "Sarry", 
            "32": "Shepherd", 
            "33": "Sapher", 
            "34": "Wingth", 
            "35": "Williams", 
            "36": "Willmore", 
            "37": "Ninguna de las anteriores" 
        },
    }
}


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def checkedElement():
    elm = OxmlElement('w:checked')
    elm.set(qn('w:val'),"true")
    return elm

zipcode_data_lookup = {}
country_index_score = {}

# with open('/home/pi/zipcode_data.txt') as json_file:
#     zipcode_data_lookup = json.load(json_file)

with open('country_json_data.json', 'r') as myfile:
    data=myfile.read()

# parse file
country_index_score = json.loads(data)

def score_answers(userInfo, country_score):
    number_correct = 0
    for answer in correct_answers:
        print(answer, userInfo[answer])
        if (userInfo[answer] == correct_answers[answer]):
            print("Correct")
            number_correct +=1

    if country_score <= 23.6:
        number_correct +=1 
    return number_correct


def formatDocument(userInfo):
    print("Starting Custom Print Job")
    print(userInfo)
    doc = Document("/Users/amilcook/Documents/printouts/CivilResponsesEN.docx")
    lang = userInfo["lang"]
    if userInfo["lang"] == "es":
        doc = Document("/Users/amilcook/Documents/printouts/CivilResponsesES.docx")


    country_name = country_index_score[userInfo["countryName"]]["country_name"]
    country_score = float(country_index_score[userInfo["countryName"]]["score"])

    print(country_name, country_score)

    styles = doc.styles
    style = styles.add_style('Insertion', WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles['Insertion']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(18)


    num_correct = score_answers(userInfo, country_score)
    qualify_status = "QUALIFY"  if num_correct == 4  else "DISQUALIFY"
    print("QUALIFY : " , qualify_status)
    print("NUM CORRECT ", num_correct)


    print("Creating Document using this info")
    print(userInfo)
    for paragraph in doc.paragraphs:
        # print(paragraph.text)

        if '[DATE]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text =f"{datetime.datetime.now():%m-%d-%Y}" 

        if '[NAME]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text = 'Applicant: {}'.format(userInfo["userName"])
            if lang == "es":
                paragraph.text = 'Solicitante: {}'.format(userInfo["userName"])

        if '[QUALIFYSTATUS]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text = 'Result : {}'.format(qualify_status)
            if lang == "es":
                paragraph.text = "Resultado : No Califica"

        if '[Q1]' in paragraph.text:   
            q1Answer = answer_lookup[lang]["a1"][userInfo["a1"]]

            if lang == "en":
                paragraph.text = "You [{}] know or have heard of the suspects and witnesses.".format(q1Answer)
            else:
                q1Answer = answer_lookup[lang]["a1"][userInfo["a1"]]
                paragraph.text = "Usted [{}] los testigos o los sospechosos.".format(q1Answer)

        if '[Qc2]' in paragraph.text:
            questionTwoAnswer = answer_lookup[lang]["a2"][userInfo["a2"]]
            paragraph.text = "You [{}] related or know a descendent of an immigrant from the 1824 migration from the United States to Haiti/Dominican Republic.".format(questionTwoAnswer)
  
            if lang == "es":
                paragraph.text = "Usted [{}] ser un familiar o conoce a un descendiente de inmigrante de la migración de 1824 de los Estados Unidos a Haití/República Dominicana.".format(questionTwoAnswer)

        if '[Q3]' in paragraph.text:
            questionThreeAnswer = answer_lookup[lang]["a3"][userInfo["a3"]]
            print("paragraph text", paragraph.text)
            paragraph.text = "For you, [{}] affects your material conditions.".format(questionThreeAnswer)
            if lang == "es":
                paragraph.text = "Para usted, [{}] afecta sus condiciones materiales.".format(questionThreeAnswer)

#        if '[Sugar]' in paragraph.text:
#            questionFourAnswer = answer_lookup[lang]["sugarIntake"][userInfo["sugarIntake"]]
#            paragraph.text = "For you, [{}] affects your material conditions.".format(answer_lookup[lang]["sugarIntake"][userInfo["sugarIntake"]])
#            if lang == "es":
#                paragraph.text = "Su consumo semanal de azúcar es [Q3 answer].".format(questionFourAnswer)

        if '[ANSWER]' in paragraph.text:
            print("Setting Random Value 1")
            paragraph.style = 'Insertion'

            paragraph.text = "Your [{}] nationality ranks {}% in the Quality of Nationality index".format(country_name,country_score)
            if lang == "es":
                paragraph.text = "Su nacionalidad [{}] tiene un índice de {}% en el índice de calidad de la nacionalidad".format(country_name,country_score)


        if '[X out of 4]' in paragraph.text or '[X de 4]' in paragraph.text:
            print("Setting Random Value 1")
            paragraph.style = 'Insertion'

            paragraph.text = "You answered [{} out of 4] questions correctly. To be an impartial reviewer you would have to answer all the questions correctly.".format(num_correct)
            if lang == "es":
                paragraph.text = "Usted obtuvo [{} de 4] correctas. Para ser un evaluador imparcial debe responder correctamente todas las preguntas.".format(num_correct)





    doc_name = '{}.docx'.format(userInfo["userName"])
    doc.save(doc_name)
    print("Formatted And Saved Document with name {}".format(doc_name))
#    subprocess.run(["libreoffice", "--headless", "--convert-to",
#                   "pdf", "{}.docx".format(userInfo["userName"])])
#    subprocess.run(
#        ["lp", "-d", "myprinter", "{}.pdf".format(userInfo["userName"])])
#    print("Completed Format Of Document")


if __name__ == "__main__":
    try:
        formatDocument({"userName":"mattest1","userId":"d30ecf6db2dc97f363a57af1bf4f4658","a1":"1","a2":"0","a3":"3","countryName":"7","sugarIntake":"1","archivePermission":"1","lang":"en"})
    except Exception as e:
        print(e)

